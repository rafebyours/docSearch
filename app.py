import os
import pickle
from flask import Flask, request, render_template, send_file, redirect, url_for, flash, session
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
from langdetect import detect
from preprocessing import preprocess_text
import PyPDF2
import docx
from fuzzywuzzy import process
import re
import csv
from concurrent.futures import ProcessPoolExecutor
from tqdm import tqdm
import sys
from flask import send_from_directory
from difflib import SequenceMatcher


app = Flask(__name__)


UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'csv'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
import os

FAVORITE_DOCUMENTS_FILE = 'favorite_documents.pkl' 



def load_documents():
    global DOCUMENTS, LANGUAGES, FILE_TITLES, FILE_TYPES, FILE_YEARS, TFIDF_MATRIX, VECTOR_MODEL

    if os.path.exists('model.pkl'):
        DOCUMENTS, LANGUAGES, TFIDF_MATRIX, VECTOR_MODEL, FILE_TITLES, FILE_TYPES, FILE_YEARS = load_model_and_data()
    else:
        documents = []
        languages = []
        file_titles = []
        file_types = []
        file_years = []
        files = os.listdir(UPLOAD_FOLDER)

        print("Reading files...")
        with ProcessPoolExecutor() as executor:
            file_contents = list(
                tqdm(
                    executor.map(read_file, [os.path.join(UPLOAD_FOLDER, f) for f in files]),
                    desc="Reading files",
                    unit="file",
                    total=len(files),
                    file=sys.stdout
                )
            )

        print("Processing files...")
        for i in tqdm(range(len(files)), desc="Processing files", unit="file", file=sys.stdout):
            file_path = files[i]
            content = file_contents[i]
            preprocessed_content = preprocess_text(content)
            documents.append(preprocessed_content)
            languages.append(detect_language(content))
            file_titles.append(file_path)

            file_type, file_year = extract_metadata(file_path)
            file_types.append(file_type)
            file_years.append(file_year)

        print("Training TF-IDF model...")
        vector_model = TfidfVectorizer()
        tfidf_matrix = vector_model.fit_transform(
            tqdm(documents, desc="Fitting TF-IDF", unit="document", file=sys.stdout)
        )

        save_model_and_data(documents, languages, tfidf_matrix, vector_model, file_titles, file_types, file_years)

        DOCUMENTS = documents
        LANGUAGES = languages
        FILE_TITLES = file_titles
        FILE_TYPES = file_types
        FILE_YEARS = file_years
        TFIDF_MATRIX = tfidf_matrix
        VECTOR_MODEL = vector_model

def load_model_and_data():
    with open('model.pkl', 'rb') as model_file:
        data = pickle.load(model_file)

    if len(data) != 7:
        raise ValueError(f"Expected 7 values, but got {len(data)}")
    return data

def read_file(file_path):
    ext = file_path.split('.')[-1]
    content = ''

    if ext == 'pdf':
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                content += page.extract_text()
    elif ext == 'docx':
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            content += paragraph.text
    elif ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    elif ext == 'csv':
        df = pd.read_csv(file_path)
        content = df.to_string(index=False)

    return content


def detect_language(text):
    try:
        lang = detect(text)
        if lang == 'id':
            return 'indonesia'
        elif lang in ['en', 'eng']:
            return 'inggris'
        else:
            return 'unknown'
    except:
        return 'unknown'

def extract_metadata(filename):
    ext = filename.split('.')[-1]
    year = re.findall(r'\d{4}', filename)
    year = year[0] if year else None
    return ext, year

def save_model_and_data(documents, languages, tfidf_matrix, vector_model, file_titles, file_types, file_years):
    with open('model.pkl', 'wb') as model_file:
        pickle.dump((documents, languages, tfidf_matrix, vector_model, file_titles, file_types, file_years), model_file)

@app.route('/', methods=['GET', 'POST'])
def home():
    load_documents()

    if request.method == 'POST':
        query = request.form.get('query', '')
        file_type_filter = request.form.get('file_type')
        language_filter = request.form.get('language')
        if not query:
            return render_template('index.html', error="Please enter a search query.")

        # Preprocessing query untuk menghapus stop words
        preprocessed_query = preprocess_text(query)
        suggested_query = suggest_query(preprocessed_query, DOCUMENTS)
        final_query = suggested_query if suggested_query else preprocessed_query
        query_vector = VECTOR_MODEL.transform([preprocessed_query])
        similarities = cosine_similarity(query_vector, TFIDF_MATRIX).flatten()
        results = []
        for i in range(len(similarities)):
            if similarities[i] > 0:
                if file_type_filter and FILE_TYPES[i] != file_type_filter:
                    continue
                if language_filter and LANGUAGES[i] != language_filter:
                    continue
                snippets = extract_snippets(DOCUMENTS[i], final_query)
                results.append({
                    'title': FILE_TITLES[i],
                    'score': similarities[i],
                    'snippets': snippets,
                    'type': FILE_TYPES[i],
                    'language': LANGUAGES[i]
                })
        results = sorted(results, key=lambda x: x['score'], reverse=True)
        if not results:
            return render_template('index.html', error="No matching documents found.", suggestion=query)

        favorite_documents = load_favorite_documents()

        return render_template('index.html', results=results, suggestion=suggested_query, favorite_documents=favorite_documents)

    favorite_documents = load_favorite_documents()
    return render_template('index.html', favorite_documents=favorite_documents)

def correct_spelling(text):
    from textblob import TextBlob
    blob = TextBlob(text)
    corrected_text = blob.correct()
    return str(corrected_text)

def suggest_query(query, documents):
    all_words = set(word for doc in documents for word in doc.split())
    suggestion, score = process.extractOne(query, all_words)
    return suggestion if score > 80 else None

import re

def highlight_words(text, words):
    highlighted = text
    for word in words:
        pattern = re.compile(r'\b' + re.escape(word) + r'\b', re.IGNORECASE)
        highlighted = pattern.sub(r'<strong>\g<0></strong>', highlighted)
    return highlighted

def extract_snippets(document, query, max_length=200):
    # Find positions of query words in the document
    query_words = query.split()
    relevant_snippets = []
    document_length = len(document)

    for word in query_words:
        # Find all positions of the query word in the document
        for match in re.finditer(r'\b' + re.escape(word) + r'\b', document, re.IGNORECASE):
            start_pos = match.start()
            end_pos = match.end()

            # Get a portion of text around the match to form the snippet
            snippet_start = max(0, start_pos - max_length // 2)  # Half before the match
            snippet_end = min(document_length, end_pos + max_length // 2)  # Half after the match

            snippet = document[snippet_start:snippet_end]

            # Highlight query words in the snippet
            snippet = highlight_words(snippet, query_words)

            # Trim snippet to ensure it's within the max_length
            if len(snippet) > max_length:
                snippet = snippet[:max_length] + '...'

            relevant_snippets.append(snippet)

            # Limit to 3 snippets per document
            if len(relevant_snippets) >= 3:
                break

        # If 3 snippets are collected, stop processing
        if len(relevant_snippets) >= 3:
            break

    return relevant_snippets

def similar_snippets(snippet1, snippet2, threshold=0.8):
    """Check if two snippets are too similar using sequence matcher"""
    s1 = re.sub(r'<[^>]+>', '', snippet1)
    s2 = re.sub(r'<[^>]+>', '', snippet2)
    return SequenceMatcher(None, s1, s2).ratio() > threshold

def save_to_favorites(filename):
    if os.path.exists(FAVORITE_DOCUMENTS_FILE):
        with open(FAVORITE_DOCUMENTS_FILE, 'rb') as f:
            favorite_documents = pickle.load(f)
    else:
        favorite_documents = []

    if filename not in favorite_documents:
        favorite_documents.append(filename)
        with open(FAVORITE_DOCUMENTS_FILE, 'wb') as f:
            pickle.dump(favorite_documents, f)
        flash(f"Dokumen {filename} berhasil ditambahkan ke favorit!", "success")
    else:
        flash(f"Dokumen {filename} sudah ada di favorit.", "info")

def load_favorite_documents():
    if os.path.exists(FAVORITE_DOCUMENTS_FILE):
        with open(FAVORITE_DOCUMENTS_FILE, 'rb') as f:
            return pickle.load(f)
    return []

def remove_from_favorites(filename):
    if os.path.exists(FAVORITE_DOCUMENTS_FILE):
        with open(FAVORITE_DOCUMENTS_FILE, 'rb') as f:
            favorite_documents = pickle.load(f)
        if filename in favorite_documents:
            favorite_documents.remove(filename)
            with open(FAVORITE_DOCUMENTS_FILE, 'wb') as f:
                pickle.dump(favorite_documents, f)
            flash(f"Dokumen {filename} berhasil dihapus dari favorit!", "success")
        else:
            flash(f"Dokumen {filename} tidak ditemukan di favorit.", "warning")

@app.route('/save/<filename>')
def save_document(filename):
    save_to_favorites(filename)
    return redirect(url_for('view_saved_documents'))

@app.route('/unfavorite/<filename>', methods=['POST'])
def unfavorite(filename):
    remove_from_favorites(filename)
    return redirect(url_for('view_saved_documents'))

@app.route('/favorites')
def view_saved_documents():
    favorite_documents = load_favorite_documents()
    return render_template('favorites.html', favorite_documents=favorite_documents)

@app.route('/download/<filename>')
def download(filename):
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    return send_file(file_path, as_attachment=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/preview/<filename>')
def preview(filename):
    file_extension = filename.split('.')[-1].lower()

    # Validasi file yang didukung
    if file_extension not in ALLOWED_EXTENSIONS:
        return "File type not allowed", 400

    # Tentukan path lengkap ke file
    file_path = os.path.join(UPLOAD_FOLDER, filename)

    # Cek apakah file ada
    if not os.path.exists(file_path):
        return "File not found", 404

    # Untuk PDF, tampilkan dengan iframe
    if file_extension == 'pdf':
        return render_template('preview.html', 
                               title=filename,  
                               file_path=f"/uploads/{filename}",  
                               download_path=f"/download/{filename}")

    # Untuk DOCX, konversi ke format HTML dan tampilkan
    elif file_extension == 'docx':
        doc = docx.Document(file_path)
        content = "" 
        for paragraph in doc.paragraphs: 
            content += f"<p>{paragraph.text}</p>"
        return render_template('preview.html',  
                               title=filename,  
                               content=content,  
                               download_path=f"/download/{filename}")

    # Untuk TXT, tampilkan sebagai teks
    elif file_extension == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return render_template('preview.html',  
                               title=filename,  
                               content=f"<pre>{content}</pre>",  
                               download_path=f"/download/{filename}")

    # Untuk CSV, baca dan tampilkan sebagai tabel
    elif file_extension == 'csv':
        with open(file_path, 'r', encoding='utf-8') as f:
            csv_reader = csv.reader(f)
            table_html = "<table border='1' cellpadding='5' cellspacing='0'><thead><tr>"
            
            # Membaca header
            headers = next(csv_reader)
            for header in headers:
                table_html += f"<th>{header}</th>"
            table_html += "</tr></thead><tbody>"

            # Membaca baris data
            for row in csv_reader:
                table_html += "<tr>"
                for column in row:
                    table_html += f"<td>{column}</td>"
                table_html += "</tr>"
            table_html += "</tbody></table>"

        return render_template('preview.html',  
                               title=filename,  
                               content=table_html,  
                               download_path=f"/download/{filename}")

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)